#!/usr/bin/env python3
"""
Reparador de plantilla Word para docxtpl
Convierte campos como {{EL O LA JUEZ}} → {{EL_O_LA_JUEZ}}
"""

import re
import os
from docx import Document

def reparar_plantilla(input_path, output_path):
    """
    Repara una plantilla Word para hacerla compatible con docxtpl
    """
    print(f"🔧 REPARANDO PLANTILLA: {input_path}")
    
    if not os.path.exists(input_path):
        print(f"❌ Error: No se encuentra archivo: {input_path}")
        return False
    
    try:
        # Cargar documento
        doc = Document(input_path)
        
        # Contadores
        campos_reparados = 0
        campos_totales = 0
        
        # Reparar en párrafos
        for para in doc.paragraphs:
            if para.text:
                original = para.text
                # Buscar y reemplazar campos {{XXX YYY ZZZ}}
                def reemplazar_campo(match):
                    campo_original = match.group(1)
                    # Convertir a nombre de variable válido
                    campo_reparado = re.sub(r'[^a-zA-Z0-9_]', '_', campo_original)
                    campo_reparado = re.sub(r'_+', '_', campo_reparado)  # Eliminar múltiples _
                    campo_reparado = campo_reparado.strip('_')
                    return '{{' + campo_reparado + '}}'
                
                texto_reparado = re.sub(r'\{\{([^}]+)\}\}', reemplazar_campo, original)
                
                if texto_reparado != original:
                    campos_en_parrafo = len(re.findall(r'\{\{([^}]+)\}\}', original))
                    campos_totales += campos_en_parrafo
                    campos_reparados += campos_en_parrafo
                    para.text = texto_reparado
        
        # Reparar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        original = cell.text
                        texto_reparado = re.sub(r'\{\{([^}]+)\}\}', 
                                              lambda m: '{{' + re.sub(r'[^a-zA-Z0-9_]', '_', m.group(1)) + '}}', 
                                              original)
                        if texto_reparado != original:
                            cell.text = texto_reparado
        
        # Guardar documento reparado
        doc.save(output_path)
        
        print(f"✅ Plantilla reparada guardada en: {output_path}")
        print(f"📊 Campos reparados: {campos_reparados} de {campos_totales}")
        
        # Mostrar cambios
        print("\n📝 EJEMPLOS DE CAMBIOS:")
        doc_reparado = Document(output_path)
        campos_encontrados = set()
        
        for para in doc_reparado.paragraphs[:10]:  # Primeros 10 párrafos
            campos = re.findall(r'\{\{([^}]+)\}\}', para.text)
            for campo in campos:
                if campo not in campos_encontrados:
                    campos_encontrados.add(campo)
                    if len(campos_encontrados) <= 10:  # Mostrar primeros 10
                        print(f"  • {campo}")
        
        print(f"\n🎯 Total campos únicos: {len(campos_encontrados)}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error reparando plantilla: {e}")
        import traceback
        traceback.print_exc()
        return False

def generar_contexto_ejemplo(plantilla_reparada_path):
    """
    Genera contexto de ejemplo para plantilla reparada
    """
    from docx import Document
    import re
    
    doc = Document(plantilla_reparada_path)
    campos = set()
    
    for para in doc.paragraphs:
        campos.update(re.findall(r'\{\{([^}]+)\}\}', para.text))
    
    # Crear contexto con valores de ejemplo
    contexto = {}
    for campo in sorted(campos):
        # Generar valor de ejemplo basado en nombre del campo
        if 'NOMBRE' in campo:
            contexto[campo] = f"EJEMPLO_{campo}"
        elif 'DNI' in campo:
            contexto[campo] = '12345678'
        elif 'EDAD' in campo:
            contexto[campo] = '30'
        elif 'FECHA' in campo:
            contexto[campo] = '19/04/2026'
        elif 'HORA' in campo:
            contexto[campo] = '09:00:00'
        else:
            contexto[campo] = f'[{campo}]'
    
    return contexto

def main():
    print("🔧 REPARADOR DE PLANTILLAS WORD PARA DOCXTPL")
    print("=" * 50)
    
    # Ruta a tu plantilla original (AJUSTA ESTA RUTA)
    plantilla_original = "A.A.A.A---46e0dd0d-4e26-47a5-aa90-d6fbe89963d7.docx"
    
    if not os.path.exists(plantilla_original):
        print(f"⚠️  No se encuentra: {plantilla_original}")
        print("\n📁 Archivos disponibles:")
        for archivo in os.listdir('.'):
            if archivo.endswith('.docx'):
                print(f"  • {archivo}")
        
        # Preguntar al usuario
        plantilla_original = input("\n📝 Ingresa nombre de plantilla a reparar: ").strip()
    
    if not os.path.exists(plantilla_original):
        print("❌ No se puede continuar sin plantilla.")
        return
    
    # Ruta para plantilla reparada
    plantilla_reparada = "plantilla_reparada.docx"
    
    # Reparar plantilla
    if reparar_plantilla(plantilla_original, plantilla_reparada):
        print(f"\n🎯 ¡PLANTILLA REPARADA EXITOSAMENTE!")
        print(f"📍 Usa: {plantilla_reparada} con docxtpl")
        
        # Generar contexto de ejemplo
        contexto = generar_contexto_ejemplo(plantilla_reparada)
        print(f"\n📋 Contexto de ejemplo ({len(contexto)} campos):")
        for i, (campo, valor) in enumerate(list(contexto.items())[:15], 1):
            print(f"  {i:2}. {campo}: {valor}")
        
        if len(contexto) > 15:
            print(f"  ... y {len(contexto) - 15} más")
        
        print("\n🚀 Ahora puedes usar la plantilla reparada con docxtpl")
    else:
        print("\n❌ Error en reparación")

if __name__ == "__main__":
    main()